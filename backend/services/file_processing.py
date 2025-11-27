"""
File processing service for extracting text from various file types.
Supports PDF, DOCX, TXT, MD, and images (with OCR).
All imports are lazy (inside functions) to avoid startup errors.
"""
import os
import io
from typing import Optional, Dict
from django.conf import settings

# Tesseract language codes for OCR
OCR_LANGUAGE_CODES = {
    'uz_cyrl': 'uzb_cyrl',  # Uzbek (Cyrillic)
    'uz_latn': 'uzb',  # Uzbek (Latin)
    'ru': 'rus',  # Russian
    'en': 'eng',  # English
}


def _check_library_available(library_name: str) -> bool:
    """Check if a library is available."""
    try:
        __import__(library_name)
        return True
    except ImportError:
        return False


def extract_text_from_file(
    file_content: bytes,
    file_name: str,
    file_type: Optional[str] = None,
    ocr_languages: Optional[list] = None
) -> Dict[str, any]:
    """
    Extract text from a file (PDF, DOCX, TXT, MD, or image).
    
    Args:
        file_content: File content as bytes
        file_name: Original file name
        file_type: File type ('pdf', 'docx', 'txt', 'md', 'image'). If None, detected from extension.
        ocr_languages: List of language codes for OCR (default: ['uzb', 'rus', 'eng'])
        
    Returns:
        dict with keys:
            - text: Extracted text
            - pages: Number of pages (for PDF/DOCX)
            - file_type: Detected file type
            - error: Error message if any
    """
    # Detect file type from extension if not provided
    if not file_type:
        extension = file_name.lower().split('.')[-1] if '.' in file_name else ''
        type_map = {
            'pdf': 'pdf',
            'docx': 'docx',
            'doc': 'docx',  # Treat .doc as .docx
            'txt': 'txt',
            'md': 'md',
            'jpg': 'image',
            'jpeg': 'image',
            'png': 'image',
            'gif': 'image',
            'bmp': 'image',
            'tiff': 'image',
        }
        file_type = type_map.get(extension, 'txt')
    
    # Default OCR languages
    if ocr_languages is None:
        ocr_languages = ['uzb', 'rus', 'eng']  # Uzbek, Russian, English
    
    try:
        if file_type == 'pdf':
            return _extract_text_from_pdf(file_content, ocr_languages)
        elif file_type == 'docx':
            return _extract_text_from_docx(file_content)
        elif file_type in ['txt', 'md']:
            return _extract_text_from_text(file_content)
        elif file_type == 'image':
            return _extract_text_from_image(file_content, ocr_languages)
        else:
            return {
                'text': '',
                'pages': 0,
                'file_type': file_type,
                'error': f'Unsupported file type: {file_type}'
            }
    except Exception as e:
        return {
            'text': '',
            'pages': 0,
            'file_type': file_type,
            'error': f'Error extracting text: {str(e)}'
        }


def _extract_text_from_pdf(file_content: bytes, ocr_languages: list) -> Dict[str, any]:
    """Extract text from PDF file. Uses OCR if text extraction fails."""
    # Lazy import PyPDF2
    try:
        import PyPDF2
    except ImportError:
        return {
            'text': '',
            'pages': 0,
            'file_type': 'pdf',
            'error': 'PyPDF2 library not installed'
        }
    
    try:
        # Try direct text extraction first
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text_parts = []
        total_pages = len(pdf_reader.pages)
        
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(f'--- Page {page_num + 1} ---\n{page_text}\n')
            except Exception:
                # If text extraction fails, try OCR for this page
                if _check_library_available('pdf2image') and _check_library_available('pytesseract'):
                    try:
                        from pdf2image import convert_from_bytes
                        import pytesseract
                        images = convert_from_bytes(file_content, first_page=page_num + 1, last_page=page_num + 1)
                        if images:
                            ocr_text = pytesseract.image_to_string(images[0], lang='+'.join(ocr_languages))
                            text_parts.append(f'--- Page {page_num + 1} (OCR) ---\n{ocr_text}\n')
                    except Exception as ocr_error:
                        text_parts.append(f'--- Page {page_num + 1} ---\n[Could not extract text: {str(ocr_error)}]\n')
                else:
                    text_parts.append(f'--- Page {page_num + 1} ---\n[OCR libraries not available]\n')
        
        full_text = '\n'.join(text_parts)
        
        return {
            'text': full_text,
            'pages': total_pages,
            'file_type': 'pdf',
        }
    except Exception as e:
        return {
            'text': '',
            'pages': 0,
            'file_type': 'pdf',
            'error': f'PDF extraction error: {str(e)}'
        }


def _extract_text_from_docx(file_content: bytes) -> Dict[str, any]:
    """Extract text from DOCX file."""
    # Lazy import python-docx
    try:
        from docx import Document as DocxDocument
    except ImportError:
        return {
            'text': '',
            'pages': 0,
            'file_type': 'docx',
            'error': 'python-docx library not installed'
        }
    
    try:
        docx_file = io.BytesIO(file_content)
        doc = DocxDocument(docx_file)
        
        # Extract text from all paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        full_text = '\n'.join(text_parts)
        
        return {
            'text': full_text,
            'pages': len(doc.paragraphs),  # Approximate pages
            'file_type': 'docx',
        }
    except Exception as e:
        return {
            'text': '',
            'pages': 0,
            'file_type': 'docx',
            'error': f'DOCX extraction error: {str(e)}'
        }


def _extract_text_from_text(file_content: bytes) -> Dict[str, any]:
    """Extract text from plain text file (TXT, MD)."""
    try:
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        text = None
        
        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        
        if text is None:
            text = file_content.decode('utf-8', errors='ignore')
        
        return {
            'text': text,
            'pages': 1,
            'file_type': 'txt',
        }
    except Exception as e:
        return {
            'text': '',
            'pages': 0,
            'file_type': 'txt',
            'error': f'Text extraction error: {str(e)}'
        }


def _extract_text_from_image(file_content: bytes, ocr_languages: list) -> Dict[str, any]:
    """Extract text from image using OCR."""
    # Lazy import Pillow and pytesseract
    try:
        from PIL import Image
    except ImportError:
        return {
            'text': '',
            'pages': 0,
            'file_type': 'image',
            'error': 'Pillow library not installed'
        }
    
    try:
        import pytesseract
    except ImportError:
        return {
            'text': '',
            'pages': 0,
            'file_type': 'image',
            'error': 'pytesseract library not installed'
        }
    
    try:
        image = Image.open(io.BytesIO(file_content))
        
        # Convert to RGB if needed
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Perform OCR
        lang_string = '+'.join(ocr_languages)
        text = pytesseract.image_to_string(image, lang=lang_string)
        
        return {
            'text': text,
            'pages': 1,
            'file_type': 'image',
        }
    except Exception as e:
        return {
            'text': '',
            'pages': 0,
            'file_type': 'image',
            'error': f'OCR error: {str(e)}'
        }
